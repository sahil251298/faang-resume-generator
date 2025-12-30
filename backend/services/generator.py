from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def generate_pdf_resume(data: dict, output_path: str):
    """
    Generates a FAANG-style PDF resume using ReportLab (Classic Serif Style).
    Reference: Single column, compact, serif typeset.
    """
    doc = SimpleDocTemplate(
        output_path,
        pagesize=LETTER,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    styles = getSampleStyleSheet()
    
    # --- Fonts & Styles ---
    # Using Standard 14 Fonts: Times-Roman, Times-Bold, Times-Italic
    
    # 1. Main Header (Name)
    name_style = ParagraphStyle(
        'Name',
        parent=styles['Normal'],
        fontName='Times-Bold',
        fontSize=24,
        alignment=TA_CENTER,
        spaceAfter=12,  # Increased Space
        textTransform='uppercase',
        textColor=colors.black
    )
    
    # 2. Subtitle / Contact
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=10,
        alignment=TA_CENTER,
        spaceBefore=6, # Added space before
        spaceAfter=12,
        textColor=colors.black
    )
    
    # 3. Section Headers
    section_header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Normal'],
        fontName='Times-Bold',
        fontSize=11,
        spaceBefore=10,
        spaceAfter=2,
        textTransform='uppercase',
        textColor=colors.black
    )
    
    # 4. Body Text
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=10.5,
        leading=13,
        alignment=TA_LEFT,
        spaceAfter=2
    )

    # 5. Bullets
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        firstLineIndent=0,
        leftIndent=12,
        bulletIndent=0,
        spaceAfter=1
    )

    story = []
    
    # --- HEADER ---
    contact_line = data.get('contact', '').strip()
    # Logic to ensure proper separation if the user provided "Name | Contact..." or just "Contact..."
    # The prompt usually asks for "Name | Phone..." but if user edits it, it might change.
    
    parts = [p.strip() for p in contact_line.split('|')]
    if len(parts) > 1:
        # Assuming first part is name
        name = parts[0]
        contact_info = " | ".join(parts[1:])
    else:
        # Fallback if no pipe separation found, or just name? 
        # Better safe: treat as Name, empty contact, OR check string length?
        # Let's assume the user put everything in one line if no pipe, so we split by newlines?
        if '\n' in contact_line:
            n_parts = contact_line.split('\n')
            name = n_parts[0]
            contact_info = " | ".join(n_parts[1:])
        else:
            name = contact_line # Risk: entire contact line becomes Name header
            contact_info = ""

    story.append(Paragraph(name, name_style))
    story.append(Paragraph(contact_info, contact_style))
    
    # Space after contact
    story.append(Spacer(1, 10))

    # --- SUMMARY (Fixed Position: 2 lines below) ---
    if data.get('summary'):
        story.append(Paragraph("SUMMARY", section_header_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceAfter=6, spaceBefore=4))
        story.append(Paragraph(data['summary'], body_style))
        story.append(Spacer(1, 12)) # slightly more space after summary section to ensure break

    # --- HELPERS ---
    def add_line():
        story.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceAfter=6, spaceBefore=4))

    # --- SECTIONS GENERATORS ---
    
    def generate_education():
        if data.get('education'):
            story.append(Paragraph("EDUCATION", section_header_style))
            add_line()
            # ... (rest of implementation unchanged)
            
            for edu_item in data['education']:
                parts = [x.strip() for x in edu_item.split(',')]
                
                if len(parts) >= 2:
                    uni_name = parts[0]
                    date = parts[-1]
                    mid = ", ".join(parts[1:-1])
                    
                    row1 = [Paragraph(f"<b>{uni_name}</b>", body_style), Paragraph(date, ParagraphStyle('Right', parent=body_style, alignment=TA_RIGHT))]
                    if mid:
                        row2 = [Paragraph(f"<i>{mid}</i>", body_style), ""]
                        data_table = [row1, row2]
                    else:
                        data_table = [row1]
                        
                    t = Table(data_table, colWidths=[5.5*inch, 2*inch])
                    t.setStyle(TableStyle([
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('LEFTPADDING', (0,0), (-1,-1), 0),
                        ('RIGHTPADDING', (0,0), (-1,-1), 0),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ]))
                    t.hAlign = 'LEFT'
                    story.append(t)
                else:
                    story.append(Paragraph(edu_item, body_style))
                
                story.append(Spacer(1, 4))

    def generate_experience():
        if data.get('experience'):
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", section_header_style))
            add_line()
            
            for exp_block in data['experience']:
                lines = exp_block.split('\n')
                header_line = lines[0]
                
                h_parts = [p.strip() for p in header_line.split('|')]
                
                if len(h_parts) >= 3:
                    company = h_parts[0]
                    role = h_parts[1]
                    date = h_parts[2]
                    
                    t_data = [[Paragraph(f"<b>{company}</b>", body_style), Paragraph(date, ParagraphStyle('R', parent=body_style, alignment=TA_RIGHT))]]
                    t_data.append([Paragraph(f"<i>{role}</i>", body_style), ""])
                    
                    t = Table(t_data, colWidths=[5.5*inch, 2*inch])
                    t.setStyle(TableStyle([
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('LEFTPADDING', (0,0), (-1,-1), 0),
                        ('RIGHTPADDING', (0,0), (-1,-1), 0),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ]))
                    t.hAlign = 'LEFT'
                    story.append(t)
                else:
                    story.append(Paragraph(f"<b>{header_line}</b>", body_style))
                
                for line in lines[1:]:
                    line = line.strip()
                    if line.startswith('•') or line.startswith('-'):
                        line = line[1:].strip()
                    if line:
                        story.append(Paragraph(f"• {line}", bullet_style))
                
                story.append(Spacer(1, 8))

    def generate_projects():
        if data.get('projects'):
            story.append(Paragraph("PROJECTS", section_header_style))
            add_line()
            
            for proj_block in data['projects']:
                lines = proj_block.split('\n')
                header_line = lines[0]
                
                story.append(Paragraph(f"<b>{header_line}</b>", body_style))
                
                for line in lines[1:]:
                    line = line.strip()
                    if line.startswith('•') or line.startswith('-'):
                        line = line[1:].strip()
                    if line:
                        story.append(Paragraph(f"• {line}", bullet_style))
                
                story.append(Spacer(1, 8))

    def generate_skills():
        if data.get('skills'):
            story.append(Paragraph("SKILLS", section_header_style))
            add_line()
            for skill_line in data['skills']:
                if ':' in skill_line:
                    cat, val = skill_line.split(':', 1)
                    p = Paragraph(f"<b>{cat}:</b> {val}", body_style)
                else:
                    p = Paragraph(skill_line, body_style)
                story.append(p)
            story.append(Spacer(1, 6))

    def generate_course_work():
        if data.get('course_work'):
            story.append(Paragraph("COURSE WORK", section_header_style))
            add_line()
            for item in data['course_work']:
                story.append(Paragraph(item, body_style))
            story.append(Spacer(1, 6))

    # Dispatch Map
    generators = {
        'education': generate_education,
        'experience': generate_experience,
        'projects': generate_projects,
        'skills': generate_skills,
        'course_work': generate_course_work
    }

    # Order processing
    order = data.get('section_order', ["education", "skills", "experience", "projects", "course_work"])
    
    for section in order:
        if section in generators:
            generators[section]()

    doc.build(story)
    return output_path

def generate_docx_resume(data: dict, output_path: str):
    """
    Generates a FAANG-style DOCX resume matching the PDF design.
    """
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)
    
    # Margins (0.5 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Helper to add section header with bottom border
    def add_section_header(title):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        
        # Add border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    # --- HEADER ---
    contact_line = data.get('contact', '').strip()
    parts = [p.strip() for p in contact_line.split('|')]
    if len(parts) > 1:
        name = parts[0]
        contact_info = " | ".join(parts[1:])
    else:
        if '\n' in contact_line:
            n_parts = contact_line.split('\n')
            name = n_parts[0]
            contact_info = " | ".join(n_parts[1:])
        else:
            name = contact_line
            contact_info = ""

    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run(name)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c = p_contact.add_run(contact_info)
    run_c.font.size = Pt(10)
    run_c.font.name = 'Times New Roman'
    p_contact.paragraph_format.space_after = Pt(10) # 10pt space

    # --- SUMMARY (Fixed Position: 2 lines ~ 20pt) ---
    if data.get('summary'):
        # Add extra spacing to simulate "2 lines below"
        # We already have space_after=10 on contact.
        # Add section header already adds space_before=12.
        # So total visual space is adequate.
        
        add_section_header("Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(data['summary'])
        run.font.name = 'Times New Roman'


    # --- GENERATORS ---
    
    def generate_education():
        if data.get('education'):
            add_section_header("Education")
            
            table = doc.add_table(rows=0, cols=2)
            table.autofit = False
            # Force table to not indent
            tblPr = table._element.tblPr
            tblInd = OxmlElement('w:tblInd')
            tblInd.set(qn('w:w'), "0")
            tblInd.set(qn('w:type'), "dxa")
            tblPr.append(tblInd)
            
            for item in data['education']:
                parts = [x.strip() for x in item.split(',')]
                uni_name = parts[0] if parts else item
                date = parts[-1] if len(parts) > 1 else ""
                mid = ", ".join(parts[1:-1]) if len(parts) > 2 else ""

                row_cells = table.add_row().cells
                # Force zero margins
                for cell in row_cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcMar = OxmlElement('w:tcMar')
                    left = OxmlElement('w:left')
                    left.set(qn('w:w'), "0")
                    left.set(qn('w:type'), "dxa")
                    tcMar.append(left)
                    tcPr.append(tcMar)

                # Cell 0: Uni Name + Degree
                p1 = row_cells[0].paragraphs[0]
                r1 = p1.add_run(uni_name)
                r1.bold = True
                r1.font.name = 'Times New Roman'
                
                # Cell 1: Date (Right Aligned)
                p2 = row_cells[1].paragraphs[0]
                r2 = p2.add_run(date)
                r2.font.name = 'Times New Roman'
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Row 2 for Degree if exists
                if mid:
                    row_cells2 = table.add_row().cells
                    for cell in row_cells2:
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcMar = OxmlElement('w:tcMar')
                        left = OxmlElement('w:left')
                        left.set(qn('w:w'), "0")
                        left.set(qn('w:type'), "dxa")
                        tcMar.append(left)
                        tcPr.append(tcMar)
                        
                    p3 = row_cells2[0].paragraphs[0]
                    r3 = p3.add_run(mid)
                    r3.italic = True
                    r3.font.name = 'Times New Roman'

    def generate_experience():
        if data.get('experience'):
            add_section_header("Professional Experience")
            
            for exp in data['experience']:
                lines = exp.split('\n')
                header_line = lines[0]
                h_parts = [p.strip() for p in header_line.split('|')]
                
                company = h_parts[0] if h_parts else header_line
                role = h_parts[1] if len(h_parts) > 1 else ""
                date = h_parts[2] if len(h_parts) > 2 else ""

                # Use a table for the header line to ensure alignment
                table = doc.add_table(rows=0, cols=2)
                table.autofit = False
                # Force table to not indent
                tblPr = table._element.tblPr
                tblInd = OxmlElement('w:tblInd')
                tblInd.set(qn('w:w'), "0")
                tblInd.set(qn('w:type'), "dxa")
                tblPr.append(tblInd)
                
                row_cells = table.add_row().cells
                
                # Zero padding cell 0
                tcPr = row_cells[0]._tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                left = OxmlElement('w:left')
                left.set(qn('w:w'), "0")
                left.set(qn('w:type'), "dxa")
                tcMar.append(left)
                tcPr.append(tcMar)

                p_comp = row_cells[0].paragraphs[0]
                r_comp = p_comp.add_run(company)
                r_comp.bold = True
                r_comp.font.name = 'Times New Roman'
                
                # Cell 1: Date
                p_date = row_cells[1].paragraphs[0]
                r_date = p_date.add_run(date)
                r_date.font.name = 'Times New Roman'
                p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Row 2: Role
                if role:
                    row_cells2 = table.add_row().cells
                    
                    tcPr2 = row_cells2[0]._tc.get_or_add_tcPr()
                    tcMar2 = OxmlElement('w:tcMar')
                    left2 = OxmlElement('w:left')
                    left2.set(qn('w:w'), "0")
                    left2.set(qn('w:type'), "dxa")
                    tcMar2.append(left2)
                    tcPr2.append(tcMar2)
                    
                    p_role = row_cells2[0].paragraphs[0]
                    r_role = p_role.add_run(role)
                    r_role.italic = True
                    r_role.font.name = 'Times New Roman'
                    
                # Bullets
                for line in lines[1:]:
                    line = line.strip()
                    if line.startswith('•') or line.startswith('-'):
                        line = line[1:].strip()
                    if line:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.2)
                        p.paragraph_format.space_after = Pt(1)
                        p.style = 'List Bullet'
                        run = p.add_run(line)
                        run.font.name = 'Times New Roman'

    def generate_projects():
        if data.get('projects'):
            add_section_header("Projects")
            for proj in data['projects']:
                lines = proj.split('\n')
                header = lines[0]
                
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(header)
                r.bold = True
                r.font.name = 'Times New Roman'
                
                for line in lines[1:]:
                    line = line.strip()
                    if line.startswith('•') or line.startswith('-'):
                        line = line[1:].strip()
                    if line:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.2)
                        p.paragraph_format.space_after = Pt(1)
                        p.style = 'List Bullet'
                        run = p.add_run(line)
                        run.font.name = 'Times New Roman'

    def generate_skills():
        if data.get('skills'):
            add_section_header("Skills")
            for skill_line in data['skills']:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                if ':' in skill_line:
                    cat, val = skill_line.split(':', 1)
                    r_cat = p.add_run(cat + ":")
                    r_cat.bold = True
                    r_cat.font.name = 'Times New Roman'
                    r_val = p.add_run(val)
                    r_val.font.name = 'Times New Roman'
                else:
                    r = p.add_run(skill_line)
                    r.font.name = 'Times New Roman'

    def generate_course_work():
        if data.get('course_work'):
            add_section_header("Course Work")
            for item in data['course_work']:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.15)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(item)
                run.font.name = 'Times New Roman'

    # Dispatch Map
    generators = {
        'education': generate_education,
        'experience': generate_experience,
        'projects': generate_projects,
        'skills': generate_skills,
        'course_work': generate_course_work
    }
    
    # Order processing
    order = data.get('section_order', ["education", "skills", "experience", "projects", "course_work"])
    
    for section in order:
        if section in generators:
            generators[section]()
            
    doc.save(output_path)
    return output_path
